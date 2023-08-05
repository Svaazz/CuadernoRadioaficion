import datetime
import re
from docx import Document
import os
import tkinter as tk
from tkinter import ttk, simpledialog

def print_table(data_list):
    formatted_data = "{:<15} {:<11} {:<10} {:<10} {:<10} {:<11} {:<8} {:<15} {:<15} {:<15}".format(
        "UBICACIÓN", "FRECUENCIA", "HORA", "FECHA", "MODO", "INTENSIDAD", "SQUELCH", "R-DCS", "R-CTCS", "T-CTCS")
    table_content = [formatted_data, "=" * 125]
    for data in data_list:
        location, frequency, time, date, mode, intensity, squelch, r_dcs, r_ctcs, t_ctcs = data
        formatted_data = "{:<15} {:<11.3f} {:<10} {:<10} {:<10} {:<11} {:<8} {:<15} {:<15} {:<15}".format(
            location, float(frequency), time, date, mode, int(intensity), int(squelch), r_dcs, r_ctcs, t_ctcs)
        table_content.append(formatted_data)
    return "\n".join(table_content)

def print_table_to_docx(doc, data_list):
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'UBICACIÓN'
    hdr_cells[1].text = 'FRECUENCIA'
    hdr_cells[2].text = 'HORA'
    hdr_cells[3].text = 'FECHA'
    hdr_cells[4].text = 'MODO'
    hdr_cells[5].text = 'INTENSIDAD'
    hdr_cells[6].text = 'SQUELCH'
    hdr_cells[7].text = 'R-DCS'
    hdr_cells[8].text = 'R-CTCS'
    hdr_cells[9].text = 'T-CTCS'

    for data in data_list:
        row_cells = table.add_row().cells
        row_cells[0].text = data[0]
        row_cells[1].text = f'{data[1]:.3f}'
        row_cells[2].text = data[2]
        row_cells[3].text = data[3]
        row_cells[4].text = data[4]
        row_cells[5].text = str(data[5])
        row_cells[6].text = str(data[6])
        row_cells[7].text = data[7]
        row_cells[8].text = data[8]
        row_cells[9].text = data[9]

def clear_screen():
    os.system('cls' if os.name == 'nt' else 'clear')

def get_unique_frequencies(data_list):
    unique_frequencies = {}
    for data in data_list:
        frequency = data[1]
        if frequency not in unique_frequencies:
            unique_frequencies[frequency] = data
    return list(unique_frequencies.values())

def get_valid_location():
    location = simpledialog.askstring("UBICACIÓN", "Introduce la ubicación (Dejar en blanco para NO DEFINIDA):")
    return "NO DEFINIDA" if location is None or location.strip() == "" else location.strip()

def get_valid_frequency():
    while True:
        try:
            frequency_input = simpledialog.askstring("FRECUENCIA", "Introduce la frecuencia:")
            if is_valid_frequency_input(frequency_input):
                frequency = float(frequency_input)
                if (136 <= frequency <= 174) or (400 <= frequency <= 520):
                    return frequency
                else:
                    print("Frecuencia fuera de rango. Introduce una frecuencia válida.")
            else:
                print("Formato de frecuencia incorrecto. Introduce una frecuencia con el formato adecuado (Ej: 142.300).")
        except ValueError:
            print("Entrada no válida. Introduce un número válido.")

def get_valid_time():
    while True:
        try:
            time_str = simpledialog.askstring("HORA", "Introduce la hora (HH:MM):")
            if time_str is None or time_str == "":
                return datetime.datetime.now().strftime("%H:%M")
            elif re.match(r'^\d{2}:\d{2}$', time_str):
                datetime.datetime.strptime(time_str, "%H:%M")
                return time_str
            else:
                print("Formato de hora incorrecto. Usa HH:MM o deja en blanco para usar la hora actual.")
        except ValueError:
            print("Formato de hora incorrecto. Usa HH:MM o deja en blanco para usar la hora actual.")

def get_valid_date():
    while True:
        try:
            date_str = simpledialog.askstring("FECHA", "Introduce la fecha (YYYY-MM-DD):")
            if date_str is None or date_str == "":
                return datetime.datetime.now().strftime("%Y-%m-%d")
            elif re.match(r'^\d{4}-\d{2}-\d{2}$', date_str):
                datetime.datetime.strptime(date_str, "%Y-%m-%d")
                return date_str
            else:
                print("Formato de fecha incorrecto. Usa YYYY-MM-DD o deja en blanco para usar la fecha actual.")
        except ValueError:
            print("Formato de fecha incorrecto. Usa YYYY-MM-DD o deja en blanco para usar la fecha actual.")

def get_valid_intensity():
    while True:
        try:
            intensity = simpledialog.askinteger("INTENSIDAD", "Introduce la intensidad (0-4):")
            if 0 <= intensity <= 4:
                return intensity
            else:
                print("Intensidad fuera de rango. Introduce un valor entre 0 y 4.")
        except ValueError:
            print("Entrada no válida. Introduce un número válido.")

def get_valid_squelch():
    while True:
        try:
            squelch = simpledialog.askinteger("SQUELCH", "Introduce el squelch (0-10):")
            if 0 <= squelch <= 10:
                return squelch
            else:
                print("Squelch fuera de rango. Introduce un valor entre 0 y 10.")
        except ValueError:
            print("Entrada no válida. Introduce un número válido.")

def save_data(location, frequency, time, date, mode, intensity, squelch, r_dcs, r_ctcs, t_ctcs):
    with open("data.txt", "a") as file:
        file.write(f"{location}\t{frequency}\t{time}\t{date}\t{mode}\t{intensity}\t{squelch}\t{r_dcs}\t{r_ctcs}\t{t_ctcs}\n")

def is_valid_frequency_input(input_str):
    return re.match(r'^\d+\.\d{3}$', input_str)

def export_data_to_docx(data_list, filename):
    doc = Document()
    doc.add_heading('Primeros Contactos', level=1)
    unique_frequencies = get_unique_frequencies(data_list)
    print_table_to_docx(doc, unique_frequencies)

    doc.add_page_break()
    doc.add_heading('Todos los Datos', level=1)
    print_table_to_docx(doc, data_list)

    doc.save(filename)
    print(f"Archivo {filename} exportado exitosamente.")

def load_data_from_file():
    data_list = []
    with open("data.txt", "r") as file:
        for line in file:
            data = line.strip().split("\t")
            while len(data) < 10:
                data.append("0")  # Sustituir valores no definidos por ceros
            data_list.append((data[0], float(data[1]), data[2], data[3], data[4], int(data[5]), int(data[6]), data[7], data[8], data[9]))
    return data_list

def display_table(data_list):
    table_window = tk.Tk()
    table_window.title("Tabla de Datos")
    table_frame = ttk.Frame(table_window)
    table_frame.pack(fill="both", expand=True)
    
    scrollbar = ttk.Scrollbar(table_frame, orient="vertical")
    tree = ttk.Treeview(table_frame, yscrollcommand=scrollbar.set)

    scrollbar.config(command=tree.yview)

    tree["columns"] = ("UBICACIÓN", "FRECUENCIA", "HORA", "FECHA", "MODO", "INTENSIDAD", "SQUELCH", "R-DCS", "R-CTCS", "T-CTCS")
    tree.column("#0", width=0, stretch="no")
    tree.column("UBICACIÓN", anchor="w", width=150)
    tree.column("FRECUENCIA", anchor="center", width=100)
    tree.column("HORA", anchor="center", width=100)
    tree.column("FECHA", anchor="center", width=100)
    tree.column("MODO", anchor="center", width=100)
    tree.column("INTENSIDAD", anchor="center", width=100)
    tree.column("SQUELCH", anchor="center", width=100)
    tree.column("R-DCS", anchor="center", width=150)
    tree.column("R-CTCS", anchor="center", width=150)
    tree.column("T-CTCS", anchor="center", width=150)

    tree.heading("#0", text="", anchor="w")
    tree.heading("UBICACIÓN", text="UBICACIÓN", anchor="w")
    tree.heading("FRECUENCIA", text="FRECUENCIA", anchor="center")
    tree.heading("HORA", text="HORA", anchor="center")
    tree.heading("FECHA", text="FECHA", anchor="center")
    tree.heading("MODO", text="MODO", anchor="center")
    tree.heading("INTENSIDAD", text="INTENSIDAD", anchor="center")
    tree.heading("SQUELCH", text="SQUELCH", anchor="center")
    tree.heading("R-DCS", text="R-DCS", anchor="center")
    tree.heading("R-CTCS", text="R-CTCS", anchor="center")
    tree.heading("T-CTCS", text="T-CTCS", anchor="center")

    for data in data_list:
        tree.insert("", "end", values=data)

    tree.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    table_window.mainloop()

def main():
    data_list = load_data_from_file()

    root = tk.Tk()
    root.title("Cuaderno de Radioafición")

    title_label = ttk.Label(root, text="Cuaderno de Radioafición", font=("Helvetica", 16))
    title_label.pack(pady=10)

    menu_frame = ttk.Frame(root)
    menu_frame.pack()

    def add_frequency():
        location = get_valid_location()
        frequency = get_valid_frequency()
        time = get_valid_time()
        date = get_valid_date()
        mode = simpledialog.askstring("MODO", "Introduce el modo (FM por defecto):") or "FM"
        intensity = get_valid_intensity()
        squelch = get_valid_squelch()
        subtones = simpledialog.askstring("Subtonos", "¿Deseas introducir subtonos? (Sí: S, No: N):")
        if subtones and subtones.lower() == "s":
            r_dcs = simpledialog.askstring("R-DCS", "Introduce R-DCS:")
            r_ctcs = simpledialog.askstring("R-CTCS", "Introduce R-CTCS:")
            t_ctcs = simpledialog.askstring("T-CTCS", "Introduce T-CTCS:")
            data = (location, frequency, time, date, mode, intensity, squelch, r_dcs, r_ctcs, t_ctcs)
            data_list.append(data)
            save_data(location, frequency, time, date, mode, intensity, squelch, r_dcs, r_ctcs, t_ctcs)
            print("Datos almacenados exitosamente.")
        elif subtones and subtones.lower() == "n" or subtones == "":
            data = (location, frequency, time, date, mode, intensity, squelch, "0", "0", "0")
            data_list.append(data)
            save_data(location, frequency, time, date, mode, intensity, squelch, "0", "0", "0")
            print("Datos almacenados exitosamente.")
        else:
            print("Opción no válida. Los datos no se almacenarán.")

    def display_table_window():
        display_table(data_list)

    def export_to_docx():
        if len(data_list) == 0:
            print("No hay datos para exportar.")
        else:
            filename = simpledialog.askstring("Exportar a DOCX", "Introduce el nombre del archivo (sin extensión .docx):")
            if filename:
                filename += '.docx'
                export_data_to_docx(data_list, filename)

    add_button = ttk.Button(menu_frame, text="Ingresar una frecuencia", command=add_frequency)
    add_button.grid(row=0, column=0, padx=10, pady=5)

    view_button = ttk.Button(menu_frame, text="Ver Tabla", command=display_table_window)
    view_button.grid(row=0, column=1, padx=10, pady=5)

    export_button = ttk.Button(menu_frame, text="Exportar a DOCX", command=export_to_docx)
    export_button.grid(row=0, column=2, padx=10, pady=5)

    exit_button = ttk.Button(menu_frame, text="Salir", command=root.destroy)
    exit_button.grid(row=0, column=3, padx=10, pady=5)

    root.mainloop()

if __name__ == "__main__":
    main()
